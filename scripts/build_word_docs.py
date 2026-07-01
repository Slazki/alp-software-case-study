from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSIONS = ROOT / "submissions"


def set_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def title(doc: Document, text: str, subtitle: str | None = None, size: float = 22) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color="0B2545")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run(subtitle)
        set_font(run2, size=11, color="555555")


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_labeled_para(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(body)


def compact_for_one_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    for style_name, size, before, after in [
        ("Heading 1", 12, 7, 3),
        ("Heading 2", 11, 6, 3),
        ("Heading 3", 10.5, 5, 2),
    ]:
        style = styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.05


def build_schema_doc() -> Path:
    doc = Document()
    configure_doc(doc)
    title(
        doc,
        "Schema Design and Data Cleaning Rationale",
        "ALP Software Case Study - Summit Ridge DataHub",
    )

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "The goal is to turn three spreadsheet-style exports into a normalized Supabase/Postgres source of truth for associations, board members, and vendors. The schema keeps operational facts queryable while preserving enough cleanup evidence for a reviewer to see what was inferred, normalized, or left unresolved."
    )

    doc.add_heading("Schema Design", level=1)
    bullet(doc, "associations stores the HOA identity, location, unit count, dues, fiscal year-end month, reserve balance, reserve-study status, and board email. Money uses numeric fields and dates use date fields rather than text.")
    bullet(doc, "board_members belongs to one association through a foreign key. Terms are modeled as dates, email can be null, and a natural-key index prevents exact duplicate member rows after cleaning.")
    bullet(doc, "vendors stores one canonical company per real vendor, including normalized phone, email, trade, COI status, and service area.")
    bullet(doc, "association_vendors models the many-to-many relationship: one vendor can serve many associations and one association can use many vendors. Vendors with no rows in this bridge are retained as prospects.")
    bullet(doc, "vendor_aliases preserves source-name variants such as Bright Path Landscaping LLC and Reliable Electrical Services after deduplication.")
    bullet(doc, "data_quality_flags captures missing, partial, conflicting, or review-worthy values so questionable data is visible instead of silently overwritten.")

    doc.add_heading("Cleaning Decisions", level=1)
    doc.add_heading("Associations", level=2)
    bullet(doc, "State values CA, ca, Calif., and California were normalized to CA.")
    bullet(doc, "Monthly dues and reserve balances were parsed into numeric values. N/A, TBD, and blanks became null; Stonebridge's $0 dues were retained and flagged for review.")
    bullet(doc, "Fiscal year-end values were stored as month numbers. December and 12 both became 12; missing fiscal year end was left null.")
    bullet(doc, "Date values with only month/year precision, such as 04/2022 and March 2024, were stored as the first day of the month with last_reserve_study_precision = month.")
    bullet(doc, "Missing unit counts for SR-02 and SR-06 were left null and flagged rather than estimated.")

    doc.add_heading("Board Members", level=2)
    bullet(doc, "The exact duplicate Alex Tanaka president row for SR-01 was removed.")
    bullet(doc, "Missing member emails and open or unknown term_end values were left null and flagged. These are operational follow-ups, not safe values to invent.")
    bullet(doc, "Board rows are kept as membership records scoped to an association rather than collapsed into a global person table, because the data does not prove that repeated names across HOAs identify the same real person.")

    doc.add_heading("Vendors", level=2)
    bullet(doc, "Vendors were deduplicated first by normalized email, then by phone/trade, then by normalized name/trade if no stronger identifier existed.")
    bullet(doc, "Bright Path Landscaping and Bright Path Landscaping LLC were merged; Coastline Plumbing duplicate rows were merged; Reliable Electric and Reliable Electrical Services were merged.")
    bullet(doc, "COI values were normalized from Yes/Y/yes and No/no. Conflicts would be flagged; a true value is retained when one source says a COI exists.")
    bullet(doc, "Prospects such as Evergreen Tree Care, Summit Roofing, GreenLeaf Pest Control, and ClearView Window Cleaning remain in vendors without association_vendors rows.")

    doc.add_heading("Regeneration and Review", level=1)
    doc.add_paragraph(
        "The SQL load file is generated from deterministic code. Re-running scripts/clean_and_seed.py refreshes cleaned CSVs, cleaned_data.json, and supabase/reset_and_seed.sql from the raw exports."
    )
    doc.add_paragraph(
        "The one-pager generator reads from Supabase when SUPABASE_DB_URL is set. It can call the OpenAI Responses API for the Summary / Areas to Watch section when OPENAI_API_KEY is provided; otherwise it uses a deterministic fallback so the report can still be tested locally."
    )

    doc.add_heading("AI-Use Transparency", level=1)
    doc.add_paragraph(
        "AI assistance was used to draft the schema approach, write repeatable scripts, and turn the cleanup choices into reviewer-facing prose. The actual data transformations are explicit in code and review flags are stored in the database for auditability."
    )

    path = SUBMISSIONS / "schema_and_cleaning_rationale.docx"
    doc.save(path)
    return path


def build_automation_doc() -> Path:
    doc = Document()
    configure_doc(doc)
    compact_for_one_page(doc)
    title(
        doc,
        "AI Automation Proposal: Maintenance Intake Triage Assistant",
        "Hypothetical next build for Summit Ridge DataHub",
        size=19,
    )

    doc.add_heading("Assumptions", level=1)
    add_labeled_para(
        doc,
        "Business context",
        "Summit Ridge earns recurring management fees and spends meaningful staff time turning owner emails, board requests, and vendor updates into tracked work.",
    )
    add_labeled_para(
        doc,
        "DataHub role",
        "Supabase is the source of truth for associations, board contacts, preferred vendors, COI status, and data-quality gaps.",
    )

    doc.add_heading("Problem", level=1)
    doc.add_paragraph(
        "Maintenance intake is repetitive and easy to misroute. A coordinator has to identify the association, infer the trade, judge urgency, check whether a vendor is approved, ask for missing details, and summarize the issue for the board or manager. Plain SQL can retrieve known facts, but it does not reliably interpret messy free-text messages or attachments."
    )

    doc.add_heading("What I Would Build", level=1)
    bullet(doc, "Input: owner emails, web-form submissions, photos, and forwarded vendor notes.")
    bullet(doc, "Output: a structured intake ticket with HOA, category/trade, urgency, missing information, suggested vendor shortlist, COI warnings, and a draft owner response.")
    bullet(doc, "AI earns its place by extracting intent from messy language, summarizing attachments, and writing a clear first draft. Rules and SQL handle deterministic checks such as vendor eligibility and missing COI.")

    doc.add_heading("Build Approach", level=1)
    bullet(doc, "Use an LLM with structured output to classify each request and produce a short, cited summary from the submitted text.")
    bullet(doc, "Query Supabase for association context, board contacts, vendor relationships, COI status, and existing data-quality flags.")
    bullet(doc, "Store the proposed ticket, model output, confidence, source message, and human approval status for auditability.")

    doc.add_heading("Risks and Guardrails", level=1)
    bullet(doc, "Low-confidence classifications, emergency language, legal issues, or missing HOA matches go to a human queue.")
    bullet(doc, "The assistant can draft messages and recommend vendors, but staff approve outbound communication and dispatch.")
    bullet(doc, "Vendor suggestions must show why they were selected and warn when COI is missing, expired, or unknown.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Prepared {date.today().isoformat()}")
    set_font(run, size=9, color="555555")

    path = SUBMISSIONS / "automation_proposal.docx"
    doc.save(path)
    return path


def main() -> None:
    SUBMISSIONS.mkdir(exist_ok=True)
    schema_path = build_schema_doc()
    automation_path = build_automation_doc()
    print(f"Wrote {schema_path}")
    print(f"Wrote {automation_path}")


if __name__ == "__main__":
    main()
